"""Funciones para leer, actualizar y generar documentos de licencia."""
from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Optional

from docx import Document
from docx.document import Document as DocumentType
from docx.table import _Cell, _Row
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .constants import (
    LABEL_TO_FIELD,
    SECTION_LABEL_TO_FIELD,
    CategoriaTipo,
    PersonaTipo,
    EQUIPMENT_FIELD_KEYS,
    TUBE_FIELD_KEYS,
)
from .text_utils import (
    apply_bold_text,
    normalize_label,
    normalize_placeholder_key,
    normalize_value,
)


@dataclass
class DocumentData:
    """Resultado del análisis de un documento fuente."""

    data: Dict[str, str]
    raw_labels: Dict[str, str]
    unmatched: Dict[str, str]
    persona: PersonaTipo | None
    categoria: CategoriaTipo | None
    equipment: List[Dict[str, str]]


@dataclass
class PlaceholderFragment:
    """Fragmento de texto que compone el reemplazo de un marcador."""

    text: str
    bold: bool = True


@dataclass
class PlaceholderContent:
    """Contenido a insertar en lugar de un marcador."""

    fragments: List[PlaceholderFragment]

    @classmethod
    def from_text(
        cls, text: str, *, bold: bool = True, uppercase: bool = True
    ) -> "PlaceholderContent":
        if uppercase:
            processed = normalize_value(text)
        else:
            processed = text.strip()
        if not processed:
            return cls([])
        return cls([PlaceholderFragment(processed, bold=bold)])


@dataclass
class RowEntry:
    """Representa una fila identificada dentro de una tabla."""

    label: str
    value: str
    label_cell: _Cell
    value_cell: _Cell
    inline: bool = False
    separator: str = ":"


def iter_paragraphs(doc: DocumentType) -> Iterable[Paragraph]:
    """Itera todos los párrafos del documento, incluyendo los de tablas."""

    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def extract_from_docx(path: Path) -> DocumentData:
    """Lee un documento de origen y obtiene los campos relevantes."""

    document = Document(str(path))
    data: Dict[str, str] = {}
    raw_labels: Dict[str, str] = {}
    unmatched: Dict[str, str] = {}
    equipment_entries: List[Dict[str, str]] = []
    current_equipment: Dict[str, str] | None = None

    def finalize_equipment() -> None:
        nonlocal current_equipment
        if current_equipment and any(current_equipment.values()):
            equipment_entries.append(current_equipment)
        current_equipment = None

    for table in document.tables:
        current_section: str | None = None
        for row in table.rows:
            section = _detect_section(row)
            if section:
                current_section = section
                if current_section != "EQUIPOS A LICENCIAR":
                    finalize_equipment()
                continue
            if (
                current_section == "EQUIPOS A LICENCIAR"
                and _is_equipment_header_row(row)
            ):
                finalize_equipment()
                current_equipment = None
                continue
            for entry in _parse_row_entries(row):
                label_norm = normalize_label(entry.label)
                raw_labels[label_norm] = entry.value
                key = _resolve_field_key(label_norm, current_section)
                if not key:
                    unmatched[entry.label] = entry.value
                    continue

                value = normalize_value(entry.value)
                in_equipment_section = current_section == "EQUIPOS A LICENCIAR"
                if key in EQUIPMENT_FIELD_KEYS and (
                    in_equipment_section or current_equipment
                ):
                    if current_equipment and any(current_equipment.values()):
                        existing = current_equipment.get(key, "")
                        if (
                            (key == "TIPO_DE_EQUIPO")
                            or (existing and value and existing != value)
                        ):
                            finalize_equipment()
                    if current_equipment is None:
                        current_equipment = {}
                    current_equipment[key] = value
                    data.setdefault(key, value)
                else:
                    data[key] = value
        finalize_equipment()

    cleaned_equipment: List[Dict[str, str]] = []
    for entry in equipment_entries:
        normalized_entry = {key: entry.get(key, "") for key in EQUIPMENT_FIELD_KEYS}
        if any(normalized_entry.values()):
            cleaned_equipment.append(normalized_entry)

    if not cleaned_equipment:
        fallback = {key: data.get(key, "") for key in EQUIPMENT_FIELD_KEYS}
        if any(fallback.values()):
            cleaned_equipment.append(fallback)

    if cleaned_equipment:
        first_equipment = cleaned_equipment[0]
        if first_equipment.get("RESOLUCION_EQUIPO") and not data.get("RESOLUCION"):
            data["RESOLUCION"] = first_equipment["RESOLUCION_EQUIPO"]
        if first_equipment.get("FECHA_RESOLUCION_EQUIPO") and not data.get("FECHA_RESOLUCION"):
            data["FECHA_RESOLUCION"] = first_equipment["FECHA_RESOLUCION_EQUIPO"]

    persona = PersonaTipo.from_text(data.get("TIPO_SOLICITANTE", ""))
    categoria = CategoriaTipo.from_text(data.get("CATEGORIA", ""))
    if categoria is None:
        categoria = CategoriaTipo.from_text(data.get("TIPO_DE_EQUIPO", ""))
    if categoria is None:
        for entry in cleaned_equipment:
            categoria = CategoriaTipo.from_text(entry.get("CATEGORIA_EQUIPO", ""))
            if categoria:
                data.setdefault("CATEGORIA", categoria.value)
                break

    return DocumentData(
        data=data,
        raw_labels=raw_labels,
        unmatched=unmatched,
        persona=persona,
        categoria=categoria,
        equipment=cleaned_equipment,
    )


def update_source_document(path: Path, updated: Dict[str, str]) -> None:
    """Sobrescribe el documento fuente con los valores corregidos."""

    document = Document(str(path))
    for table in document.tables:
        current_section: str | None = None
        for row in table.rows:
            section = _detect_section(row)
            if section:
                current_section = section
                continue
            for entry in _parse_row_entries(row):
                label_norm = normalize_label(entry.label)
                key = _resolve_field_key(label_norm, current_section)
                if not key:
                    continue
                value = updated.get(key)
                if value is None:
                    continue
                if entry.inline:
                    write_inline_cell(
                        entry.value_cell, entry.label, value, entry.separator
                    )
                else:
                    write_cell(entry.value_cell, value)
    document.save(str(path))


def write_cell(cell: _Cell, value: str) -> None:
    """Escribe el valor en una celda asegurando formato en mayúsculas y negrilla."""

    while cell.paragraphs:
        cell._element.remove(cell.paragraphs[0]._p)  # type: ignore[attr-defined]
    paragraph = cell.add_paragraph()
    run = paragraph.add_run()
    apply_bold_text(run, normalize_value(value))


def write_inline_cell(cell: _Cell, label: str, value: str, separator: str) -> None:
    """Actualiza celdas que contienen etiqueta y valor en el mismo bloque."""

    while cell.paragraphs:
        cell._element.remove(cell.paragraphs[0]._p)  # type: ignore[attr-defined]

    normalized_label = normalize_value(label)
    normalized_value = normalize_value(value)

    if separator == "\n":
        label_paragraph = cell.add_paragraph()
        label_run = label_paragraph.add_run(normalized_label)
        label_run.bold = True
        value_paragraph = cell.add_paragraph()
        value_run = value_paragraph.add_run()
        apply_bold_text(value_run, normalized_value)
        return

    paragraph = cell.add_paragraph()
    label_run = paragraph.add_run(normalized_label)
    label_run.bold = True
    separator_run = paragraph.add_run(f"{separator} ")
    separator_run.bold = True
    value_run = paragraph.add_run()
    apply_bold_text(value_run, normalized_value)


def replace_placeholders(
    document: DocumentType, data: Dict[str, PlaceholderContent]
) -> None:
    """Reemplaza cada marcador `{{CLAVE}}` por su valor correspondiente."""

    placeholders = {
        normalize_placeholder_key(key): value for key, value in data.items()
    }
    if not placeholders:
        return
    for paragraph in iter_paragraphs(document):
        replace_in_paragraph(paragraph, placeholders)


_PLACEHOLDER_PATTERN = re.compile(r"{{\s*([^{}]+?)\s*}}")


def replace_in_paragraph(
    paragraph: Paragraph, placeholders: Dict[str, PlaceholderContent]
) -> None:
    """Reemplaza marcadores dentro de un párrafo conservando formato."""

    if not paragraph.runs:
        return

    runs = [run for run in paragraph.runs]
    full_text = "".join(run.text for run in runs)
    if not full_text:
        return

    matches = list(_PLACEHOLDER_PATTERN.finditer(full_text))
    if not matches:
        return

    run_spans: List[tuple[int, int, Run]] = []
    index = 0
    for run in runs:
        length = len(run.text)
        run_spans.append((index, index + length, run))
        index += length

    fragments: List[tuple[str, object, int]] = []
    cursor = 0
    for match in matches:
        if match.start() > cursor:
            fragments.append(("text", full_text[cursor:match.start()], cursor))
        key = normalize_placeholder_key(match.group(1))
        content = placeholders.get(key)
        if content is None:
            fragments.append(("text", match.group(0), match.start()))
        else:
            fragments.append(("placeholder", content, match.start()))
        cursor = match.end()
    if cursor < len(full_text):
        fragments.append(("text", full_text[cursor:], cursor))

    while paragraph.runs:
        paragraph._element.remove(paragraph.runs[0]._r)

    for kind, payload, start in fragments:
        if kind == "text":
            _write_text_fragment(paragraph, str(payload), start, run_spans)
        else:
            _write_placeholder_fragment(
                paragraph, payload, start, run_spans
            )


def _find_reference_run(start: int, spans: List[tuple[int, int, Run]]) -> Optional[Run]:
    for begin, end, run in spans:
        if begin <= start < end:
            return run
    if spans:
        return spans[-1][2]
    return None


def _copy_run_format(source: Run, target: Run) -> None:
    target.style = source.style
    target.bold = source.bold
    target.italic = source.italic
    target.underline = source.underline
    if source.font is not None:
        target.font.name = source.font.name
        target.font.size = source.font.size


def _find_span_covering(
    position: int, spans: List[tuple[int, int, Run]]
) -> Optional[tuple[int, int, Run]]:
    for begin, end, run in spans:
        if begin <= position < end:
            return begin, end, run
    if spans:
        return spans[-1]
    return None


def _write_text_fragment(
    paragraph: Paragraph, text: str, start: int, spans: List[tuple[int, int, Run]]
) -> None:
    if not text:
        return
    cursor = start
    consumed = 0
    total = len(text)
    while consumed < total:
        span = _find_span_covering(cursor, spans)
        if span is None:
            run = paragraph.add_run(text[consumed:])
            run.bold = False
            break
        begin, end, reference = span
        span_limit = max(cursor, begin)
        available = end - span_limit
        if available <= 0:
            next_index = spans.index(span) + 1
            if next_index >= len(spans):
                break
            cursor = spans[next_index][0]
            continue
        take = min(available, total - consumed)
        chunk = text[consumed : consumed + take]
        run = paragraph.add_run(chunk)
        _copy_run_format(reference, run)
        consumed += take
        cursor = span_limit + take


def _write_placeholder_fragment(
    paragraph: Paragraph,
    content: PlaceholderContent,
    start: int,
    spans: List[tuple[int, int, Run]],
) -> None:
    if not content.fragments:
        return
    reference = _find_reference_run(start, spans)
    for fragment in content.fragments:
        if not fragment.text:
            continue
        run = paragraph.add_run(fragment.text)
        if reference:
            _copy_run_format(reference, run)
        run.bold = fragment.bold


_RESOLUTION_PLACEHOLDER_KEYS = {
    normalize_placeholder_key("RESOLUCION"),
    normalize_placeholder_key("DIA_EMISION"),
    normalize_placeholder_key("MES_EMISION"),
    normalize_placeholder_key("ANO_EMISION"),
    normalize_placeholder_key("DIA"),
    normalize_placeholder_key("MES"),
    normalize_placeholder_key("ANO"),
    normalize_placeholder_key("PARRAFO_RESOLUCION"),
}


def generate_from_template(
    template_path: Path,
    output_path: Path,
    data: Dict[str, str],
    *,
    equipment_entries: Optional[List[Dict[str, str]]] = None,
    include_resolution_paragraph: bool = True,
) -> Path:
    """Crea un documento a partir de la plantilla y lo guarda."""

    document = Document(str(template_path))

    working_data = dict(data)
    working_data.setdefault("DATOS_TUBO", "")
    working_data.setdefault("LISTA_EQUIPOS", "")

    normalized_equipment, base_equipment = _prepare_equipment_entries(
        working_data, equipment_entries
    )
    if base_equipment:
        working_data["DATOS_TUBO"] = _compose_tube_summary(base_equipment)

    resolution_content: PlaceholderContent | None = None
    if include_resolution_paragraph:
        resolution_content = _build_resolution_paragraph(working_data)
        if resolution_content is None:
            include_resolution_paragraph = False

    if not include_resolution_paragraph:
        _remove_resolution_paragraph(document)
        working_data["PARRAFO_RESOLUCION"] = ""
        resolution_content = None
    else:
        working_data.setdefault("PARRAFO_RESOLUCION", "")

    equipment_blocks = _build_equipment_blocks(normalized_equipment, working_data)
    if equipment_blocks:
        _inject_equipment_list(document, equipment_blocks)

    expanded = _expand_placeholder_aliases(working_data)
    placeholder_contents = _build_placeholder_contents(expanded)
    if resolution_content is not None:
        placeholder_contents["PARRAFO_RESOLUCION"] = resolution_content
    if equipment_blocks:
        placeholder_contents.pop("LISTA_EQUIPOS", None)
    replace_placeholders(document, placeholder_contents)
    document.save(str(output_path))
    return output_path


def _build_resolution_paragraph(data: Dict[str, str]) -> PlaceholderContent | None:
    """Construye el contenido del párrafo que deja sin efecto la resolución previa."""

    required = {
        "RESOLUCION": "",
        "DIA_EMISION": "",
        "MES_EMISION": "",
        "ANO_EMISION": "",
    }

    for key in required:
        value = data.get(key, "")
        if not value:
            return None
        required[key] = value.strip()

    month_text = required["MES_EMISION"].strip()
    month_sentence = month_text.lower()
    highlight = (
        f"Resolución No {required['RESOLUCION']} del {required['DIA_EMISION']} "
        f"de {month_sentence} de {required['ANO_EMISION']}"
    )
    intro = "Este acto administrativo deja sin efecto la "
    ending = (
        ", mediante la cual se había concedido licencia de práctica médica para "
        "este equipo de Rayos X."
    )

    return PlaceholderContent(
        [
            PlaceholderFragment(intro, bold=False),
            PlaceholderFragment(highlight, bold=True),
            PlaceholderFragment(ending, bold=False),
        ]
    )


def _expand_placeholder_aliases(data: Dict[str, str]) -> Dict[str, str]:
    """Agrega claves alternativas para los marcadores más comunes."""

    expanded = dict(data)

    day = data.get("DIA_EMISION")
    if day and "DIA" not in expanded:
        expanded["DIA"] = day

    month = data.get("MES_EMISION")
    if month and "MES" not in expanded:
        expanded["MES"] = month

    year = data.get("ANO_EMISION") or data.get("AÑO_EMISION")
    if year:
        expanded.setdefault("AÑO", year)
        expanded.setdefault("ANO", year)

    return expanded


def _build_placeholder_contents(data: Dict[str, str]) -> Dict[str, PlaceholderContent]:
    """Convierte los valores en contenido listo para ser reemplazado."""

    contents: Dict[str, PlaceholderContent] = {}
    for key, value in data.items():
        if key == "FECHA_HOY":
            contents[key] = PlaceholderContent.from_text(value, bold=False)
        else:
            contents[key] = PlaceholderContent.from_text(value)
    return contents


def _prepare_equipment_entries(
    data: Dict[str, str],
    equipment_entries: Optional[List[Dict[str, str]]],
) -> tuple[List[Dict[str, str]], Dict[str, str] | None]:
    """Normaliza la lista de equipos y obtiene el primero como referencia."""

    normalized: List[Dict[str, str]] = []
    if equipment_entries:
        for entry in equipment_entries:
            normalized.append(
                {key: normalize_value(entry.get(key, "")) for key in EQUIPMENT_FIELD_KEYS}
            )

    fallback = {key: normalize_value(data.get(key, "")) for key in EQUIPMENT_FIELD_KEYS}
    base = None

    if normalized:
        base = normalized[0]
    elif any(fallback.values()):
        base = fallback
    else:
        base = None

    if not normalized and base:
        normalized.append(base)

    return normalized, base


def _compose_tube_summary(entry: Dict[str, str]) -> str:
    """Construye el texto para el marcador DATOS_TUBO."""

    values = [entry.get(key, "") for key in TUBE_FIELD_KEYS]
    if values and all(not value or value == "NO REGISTRA" for value in values):
        return ""

    parts = []
    if entry.get("MARCA_TUBO"):
        parts.append(f"MARCA: {entry['MARCA_TUBO']}")
    if entry.get("MODELO_TUBO"):
        parts.append(f"MODELO: {entry['MODELO_TUBO']}")
    if entry.get("SERIE_TUBO"):
        parts.append(f"NUMERO DE SERIE: {entry['SERIE_TUBO']}")
    return " ".join(parts)


def _build_equipment_blocks(
    equipment_entries: List[Dict[str, str]],
    data: Dict[str, str],
) -> List[List[str]]:
    """Genera los párrafos que describen cada equipo."""

    blocks: List[List[str]] = []
    for index, entry in enumerate(equipment_entries, start=1):
        lines = _build_equipment_lines(entry, index, data)
        if lines:
            blocks.append(lines)
    return blocks


def _build_equipment_lines(
    entry: Dict[str, str],
    index: int,
    data: Dict[str, str],
) -> List[str]:
    """Crea las líneas descriptivas para un equipo individual."""

    category = entry.get("CATEGORIA_EQUIPO") or data.get("CATEGORIA", "")
    header_parts = [f"{index}. EQUIPO DE RAYOS X PARA PRACTICA MEDICA"]
    if category:
        header_parts.append(category)
    if entry.get("PRACTICA"):
        header_parts.append(entry["PRACTICA"])
    if entry.get("TIPO_DE_EQUIPO"):
        header_parts.append(entry["TIPO_DE_EQUIPO"])
    header = " ".join(part for part in header_parts if part).strip()

    details_segments = [
        _format_segment("MARCA", entry.get("MARCA", "")),
        _format_segment("MODELO", entry.get("MODELO", "")),
        _format_segment("NUMERO DE SERIE", entry.get("SERIE", "")),
        _format_segment("FECHA FABRICACION", entry.get("FECHA_FABRICACION", ""), "."),
    ]
    details = " ".join(segment for segment in details_segments if segment).strip()
    if details and not details.endswith("."):
        details = f"{details}."

    tube_segments: List[str] = []
    tube_summary = _compose_tube_summary(entry)
    if tube_summary:
        tube_segments.append(tube_summary)
    if entry.get("KV"):
        tube_segments.append(f"POTENCIA OPERACION: {entry['KV']} KV")
    if entry.get("MA"):
        tube_segments.append(f"CORRIENTE OPERACION: {entry['MA']} MA")
    if entry.get("FECHA_FABRICACION_TUBO"):
        tube_segments.append(f"FECHA FABRICACION: {entry['FECHA_FABRICACION_TUBO']}")
    if entry.get("W"):
        tube_segments.append(f"CARGA TRABAJO: W(MA.MIN/SEM){entry['W']}")
    tube_line = ""
    if tube_segments:
        tube_line = "TUBO DE RAYOS X " + " ".join(tube_segments)
        if not tube_line.endswith("."):
            tube_line += "."

    location = ""
    if entry.get("UBICACION_EQUIPO"):
        location = (
            "UBICACION EQUIPO RAYOS X DENTRO DE LA INSTALACION: "
            f"{entry['UBICACION_EQUIPO']}"
        )

    empresa_qc = entry.get("EMPRESA_QC", "")
    fecha_qc = entry.get("FECHA_QC", "")
    if empresa_qc and fecha_qc:
        qc_line = f"CONTROL DE CALIDAD REALIZADO POR: {empresa_qc} EL {fecha_qc}"
    elif empresa_qc:
        qc_line = f"CONTROL DE CALIDAD REALIZADO POR: {empresa_qc}"
    elif fecha_qc:
        qc_line = f"CONTROL DE CALIDAD REALIZADO POR: EL {fecha_qc}"
    else:
        qc_line = ""

    lines = [normalize_value(text) for text in (header, details, tube_line, location, qc_line) if text]
    return lines


def _format_segment(label: str, value: str, suffix: str = "") -> str:
    if not value:
        return ""
    segment = f"{label}: {value}"
    if suffix and not segment.endswith(suffix):
        segment += suffix
    return segment


def _inject_equipment_list(
    document: DocumentType, equipment_blocks: List[List[str]]
) -> None:
    """Inserta los párrafos de equipos en lugar del marcador LISTA_EQUIPOS."""

    if not equipment_blocks:
        return

    placeholders: List[Paragraph] = []
    for paragraph in iter_paragraphs(document):
        if not paragraph.text:
            continue
        matches = _PLACEHOLDER_PATTERN.findall(paragraph.text)
        if any(normalize_placeholder_key(match) == "LISTA_EQUIPOS" for match in matches):
            placeholders.append(paragraph)

    if not placeholders:
        return

    for paragraph in placeholders:
        for block in equipment_blocks:
            for line in block:
                new_paragraph = paragraph.insert_paragraph_before("")
                run = new_paragraph.add_run(normalize_value(line))
                run.bold = True
        _remove_paragraph(paragraph)


def build_output_name(
    source_file: Path, radicado: str, *, suffix: str | None = None
) -> str:
    """Genera el nombre final del archivo de licencia."""

    base = source_file.stem
    parts = base.split("_")
    if len(parts) >= 3:
        parts[-1] = "LICENCIA"
        new_name = "_".join(parts)
    else:
        new_name = f"{base}_LICENCIA"

    cleaned_suffix = ""
    if suffix:
        suffix_key = normalize_placeholder_key(suffix)
        if suffix_key:
            cleaned_suffix = f"_{suffix_key}"

    return f"{normalize_value(radicado)}_{new_name.split('_', 1)[-1]}{cleaned_suffix}"


def _parse_row_entries(row: _Row) -> List[RowEntry]:
    """Descompone una fila en una o varias parejas etiqueta-valor."""

    cells = list(row.cells)
    if not cells:
        return []

    non_empty = [
        (idx, cell.text.strip(), cell)
        for idx, cell in enumerate(cells)
        if cell.text and cell.text.strip()
    ]
    if not non_empty:
        return []

    entries: List[RowEntry] = []
    total = len(non_empty)
    i = 0

    while i < total:
        idx, text, cell = non_empty[i]
        inline = _split_inline_cell(text)
        if inline:
            label_text, value_text, separator = inline
            entries.append(
                RowEntry(
                    label=label_text,
                    value=value_text,
                    label_cell=cell,
                    value_cell=cell,
                    inline=True,
                    separator=separator,
                )
            )
            i += 1
            continue

        label_norm = normalize_label(text)
        if not _looks_like_label(text, label_norm):
            i += 1
            continue

        value_text: Optional[str] = None
        value_cell = cell

        j = i + 1
        while j < total:
            _, candidate_text, candidate_cell = non_empty[j]
            candidate_inline = _split_inline_cell(candidate_text)
            candidate_norm = normalize_label(candidate_text)

            if candidate_inline:
                break

            if _looks_like_label(candidate_text, candidate_norm):
                break

            value_text = candidate_text
            value_cell = candidate_cell
            break

        if value_text:
            entries.append(
                RowEntry(
                    label=text,
                    value=value_text,
                    label_cell=cell,
                    value_cell=value_cell,
                    inline=False,
                    separator=":",
                )
            )
            i = j
        else:
            i += 1

    return entries


def _remove_resolution_paragraph(document: DocumentType) -> None:
    """Elimina el párrafo asociado a la resolución previa si está presente."""

    paragraphs = list(iter_paragraphs(document))
    for paragraph in paragraphs:
        text = paragraph.text or ""
        if not text:
            continue
        matches = _PLACEHOLDER_PATTERN.findall(text)
        if not matches:
            continue
        normalized = {normalize_placeholder_key(match) for match in matches}
        if normalized & _RESOLUTION_PLACEHOLDER_KEYS:
            _remove_paragraph(paragraph)


def _remove_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is None:
        return
    parent.remove(paragraph._element)


def _detect_section(row: _Row) -> Optional[str]:
    """Detecta si la fila representa un encabezado de sección."""

    texts = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
    if not texts:
        return None

    dedup_texts: List[str] = []
    seen_norm: List[str] = []
    for text in texts:
        normalized_text = normalize_label(text)
        if not normalized_text:
            continue
        if seen_norm and normalized_text == seen_norm[-1]:
            continue
        seen_norm.append(normalized_text)
        dedup_texts.append(text)

    if not dedup_texts:
        return None

    candidate = " ".join(dedup_texts)
    normalized = normalize_label(candidate)
    if normalized in SECTION_LABEL_TO_FIELD:
        return normalized

    if len(dedup_texts) != 1:
        return None
    normalized = normalize_label(dedup_texts[0])
    if normalized in SECTION_LABEL_TO_FIELD:
        return normalized
    return None


def _resolve_field_key(label_norm: str, section: Optional[str]) -> Optional[str]:
    """Obtiene la clave asociada a una etiqueta considerando la sección."""

    if section:
        mapping = SECTION_LABEL_TO_FIELD.get(section, {})
        key = mapping.get(label_norm)
        if key:
            return key
    key = LABEL_TO_FIELD.get(label_norm)
    if key:
        return key
    return None


_LABEL_HINT_KEYWORDS = (
    "RADIC",
    "FECHA",
    "CATEG",
    "RESOL",
    "SOLICIT",
    "REPRESENT",
    "NIT",
    "CEDULA",
    "DIREC",
    "MUNIC",
    "SUBREG",
    "SEDE",
    "EQUIPO",
    "MARCA",
    "MODELO",
    "SERIE",
    "TUBO",
    "CONTROL",
    "EMPRES",
    "QC",
    "OPR",
    "OBSERV",
)


def _looks_like_label(text: str, normalized: str) -> bool:
    """Determina si un bloque de texto se comporta como etiqueta."""

    if not text:
        return False

    if normalized in LABEL_TO_FIELD:
        return True

    stripped = text.strip()
    if not stripped:
        return False

    if ":" in stripped:
        return True

    for hint in _LABEL_HINT_KEYWORDS:
        if hint in normalized:
            return True

    return False


_EQUIPMENT_HEADER_RE = re.compile(
    r"^EQUIPO(?:\s+(?:NO\.?|N[º°])\s*)?(?:\s*\d+)?[\s:.-]*$"
)


def _is_equipment_header_row(row: _Row) -> bool:
    """Determina si una fila corresponde al encabezado de un equipo."""

    texts = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
    if not texts:
        return False

    dedup_texts: List[str] = []
    seen_norm: List[str] = []
    for text in texts:
        normalized_text = normalize_label(text)
        if not normalized_text:
            continue
        if seen_norm and normalized_text == seen_norm[-1]:
            continue
        seen_norm.append(normalized_text)
        dedup_texts.append(text)

    if not dedup_texts:
        return False

    candidate = " ".join(dedup_texts)
    normalized = normalize_label(candidate)
    if not normalized:
        return False

    if normalized.startswith("EQUIPO"):
        if "A LICENCIAR" in normalized or "TIPO" in normalized:
            return False
        return bool(_EQUIPMENT_HEADER_RE.match(normalized))

    return False


def _split_inline_cell(text: str) -> Optional[tuple[str, str, str]]:
    """Divide el contenido de una celda etiqueta:valor en línea."""

    if not text:
        return None
    stripped = text.strip()
    if not stripped:
        return None

    for separator in (":", "–", "—", ";"):
        if separator in stripped:
            left, right = stripped.split(separator, 1)
            label = left.strip()
            value = right.strip()
            if label and value and _looks_like_label(label, normalize_label(label)):
                return label, value, separator

    if "-" in stripped:
        left, right = stripped.split("-", 1)
        label = left.strip()
        value = right.strip()
        if (
            label
            and value
            and _looks_like_label(label, normalize_label(label))
            and not value.replace("-", "").isdigit()
        ):
            return label, value, "-"

    lines = [line.strip() for line in stripped.splitlines() if line.strip()]
    if len(lines) >= 2:
        label = lines[0]
        value = " ".join(lines[1:])
        if label and value:
            return label, value, "\n"

    return None
